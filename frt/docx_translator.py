#!/usr/bin/env python3
"""
DOCX Format-Preserving Translation Script
This script extracts text from DOCX files while preserving formatting,
translates the text, and re-applies the formatting to the translated text.
"""

import sys
import json
import os
import requests
from docx import Document
from docx.text.run import Run
from docx.text.paragraph import Paragraph


def extract_text_with_formatting(docx_path):
    """
    Extract text from DOCX file while preserving formatting information
    """
    doc = Document(docx_path)
    content = []
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_data = {
            'text': paragraph.text,
            'style': paragraph.style.name if paragraph.style else None,
            'alignment': paragraph.alignment,
            'runs': []
        }
        
        # Extract formatting for each run in the paragraph
        for run_idx, run in enumerate(paragraph.runs):
            run_data = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
                'highlight': run.font.highlight_color
            }
            para_data['runs'].append(run_data)
        
        content.append(para_data)
    
    return content


def apply_translated_text_to_docx(original_docx_path, translated_content, output_path):
    """
    Apply translated text to a new DOCX file while preserving original formatting
    """
    doc = Document(original_docx_path)
    
    for para_idx, (paragraph, para_data) in enumerate(zip(doc.paragraphs, translated_content)):
        if para_idx >= len(translated_content):
            break
            
        translated_para = translated_content[para_idx]
        translated_text = translated_para.get('text', '')
        
        # Clear existing runs
        paragraph.clear()
        
        # Recreate runs with translated text but original formatting
        if translated_para.get('runs'):
            # Distribute translated text among runs
            total_chars = len(translated_text)
            num_runs = len(translated_para['runs'])
            
            if num_runs > 0:
                chars_per_run = total_chars // num_runs
                extra_chars = total_chars % num_runs
                
                start_idx = 0
                for run_idx, run_data in enumerate(translated_para['runs']):
                    # Calculate how many characters this run should get
                    chars_for_this_run = chars_per_run + (1 if run_idx < extra_chars else 0)
                    end_idx = start_idx + chars_for_this_run
                    
                    if end_idx > total_chars:
                        end_idx = total_chars
                        
                    # Extract the portion of translated text for this run
                    run_text = translated_text[start_idx:end_idx] if start_idx < total_chars else ""
                    start_idx = end_idx
                    
                    # Create new run with original formatting
                    new_run = paragraph.add_run(run_text)
                    new_run.bold = run_data.get('bold', None)
                    new_run.italic = run_data.get('italic', None)
                    new_run.underline = run_data.get('underline', None)
                    
                    # Apply font properties if they exist
                    if run_data.get('font_name'):
                        new_run.font.name = run_data['font_name']
                    if run_data.get('font_size'):
                        new_run.font.size = run_data['font_size']
                    if run_data.get('color'):
                        new_run.font.color.rgb = run_data['color']
                    if run_data.get('highlight'):
                        new_run.font.highlight_color = run_data['highlight']
        else:
            # If no run data, just add the text as a single run
            paragraph.add_run(translated_text)
    
    # Save the document
    doc.save(output_path)
    return output_path


def translate_text_with_llm(text, source_lang, target_lang, api_key=None):
    """
    Translate text using an LLM API
    """
    if not text.strip():
        return text
    
    # This would call your actual translation service
    # For now, we'll use a mock implementation
    try:
        # In a real implementation, you would call your LLM API here
        # Example with OpenAI-style API:
        # response = requests.post(
        #     "https://api.openai.com/v1/chat/completions",
        #     headers={
        #         "Authorization": f"Bearer {api_key}",
        #         "Content-Type": "application/json"
        #     },
        #     json={
        #         "model": "gpt-3.5-turbo",
        #         "messages": [{
        #             "role": "user",
        #             "content": f"Translate the following text from {source_lang} to {target_lang}. Provide only the translation without any additional text or explanations:\n\nText: {text}\n\nTranslation:"
        #         }],
        #         "temperature": 0.3
        #     }
        # )
        # result = response.json()
        # return result['choices'][0]['message']['content'].strip()
        
        # For demonstration, we'll just return a mock translation
        return f"[Translated] {text}"
    except Exception as e:
        print(f"Translation error: {e}", file=sys.stderr)
        return text  # Return original text if translation fails


def main():
    if len(sys.argv) < 4:
        print("Usage: python3 docx_translator.py <input_docx> <source_lang> <target_lang> [output_docx]")
        sys.exit(1)
    
    input_docx = sys.argv[1]
    source_lang = sys.argv[2]
    target_lang = sys.argv[3]
    output_docx = sys.argv[4] if len(sys.argv) > 4 else "translated_" + os.path.basename(input_docx)
    
    if not os.path.exists(input_docx):
        print(f"Error: Input file {input_docx} not found")
        sys.exit(1)
    
    try:
        # Extract text with formatting
        content = extract_text_with_formatting(input_docx)
        
        # Translate each paragraph
        translated_content = []
        for para_data in content:
            translated_text = translate_text_with_llm(para_data['text'], source_lang, target_lang)
            para_data['text'] = translated_text
            translated_content.append(para_data)
        
        # Apply translated text to new DOCX file
        result_path = apply_translated_text_to_docx(input_docx, translated_content, output_docx)
        
        # Output result information as JSON
        result = {
            "success": True,
            "input_file": input_docx,
            "output_file": result_path,
            "source_language": source_lang,
            "target_language": target_lang
        }
        print(json.dumps(result))
        
    except Exception as e:
        error_result = {
            "success": False,
            "error": str(e),
            "input_file": input_docx
        }
        print(json.dumps(error_result))
        sys.exit(1)


if __name__ == "__main__":
    main()